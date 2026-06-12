"""合同格式规范化工具 - LLM 驱动的智能格式复制

核心策略：
1. 从参考文档提取：页边距、编号定义、页眉页脚（含页码）、字体样式
2. 用 LLM 智能分析段落结构（分类 + 手动前缀检测），支持多轮验证
3. 应用参考文档的编号定义和格式到测试文档
4. LLM 失败时 fallback 到简单关键词规则
"""

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocxFormatNormalizer:  # pragma: no cover
    """合同格式规范化器 - LLM 驱动"""

    def __init__(  # pragma: no cover
        self,
        input_path: str | Path,
        output_path: str | Path | None = None,
        reference_path: str | Path | None = None,
    ):
        self.input_path = Path(input_path)
        self.output_path = (
            Path(output_path)
            if output_path
            else self.input_path.parent / f"{self.input_path.stem}_规范化{self.input_path.suffix}"
        )
        self.reference_path = Path(reference_path) if reference_path else None
        self.doc: Any = None
        self.ref_doc: Any = None
        # LLM 分析结果缓存：{para_index: {"level": int, "prefix": str}}
        self._llm_results: dict[int, dict[str, Any]] = {}

    def normalize(self, use_llm: bool = True, llm_backend: str = "openai_compatible") -> Path:  # pragma: no cover
        """执行格式规范化，返回输出文件路径"""
        logger.info("开始规范化: %s (use_llm=%s)", self.input_path, use_llm)
        self.doc = Document(str(self.input_path))

        # 用 LLM 预分析所有段落（一次性批量调用）
        if use_llm:
            self._llm_results = self._llm_analyze_document(llm_backend)

        # 加载参考文档
        if self.reference_path and self.reference_path.exists():
            self.ref_doc = Document(str(self.reference_path))
            logger.info("已加载参考文档: %s", self.reference_path)
            self._normalize_with_reference()
        else:
            logger.info("无参考文档，使用内置默认格式")
            self._normalize_default()

        assert self.doc is not None
        self.doc.save(str(self.output_path))
        logger.info("规范化完成: %s", self.output_path)
        return self.output_path

    # ── LLM 批量分析 ────────────────────────────────────────

    def _llm_analyze_document(self, llm_backend: str) -> dict[int, dict[str, Any]]:  # pragma: no cover
        """用 LLM 分析整个文档，返回每个段落的分类和前缀信息"""
        assert self.doc is not None

        # 收集非空段落
        paragraphs = []
        para_indices = []
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                para_indices.append(i)

        if not paragraphs:
            return {}

        try:
            from apps.contract_review.services.format_normalizer.llm_helper import ContractStructureAnalyzer

            analyzer = ContractStructureAnalyzer()  # 自动选择可用后端
            llm_results = analyzer.analyze_document(paragraphs, max_rounds=2)

            # 映射回段落索引
            results: dict[int, dict[str, Any]] = {}
            for idx, result in zip(para_indices, llm_results):
                results[idx] = {
                    "level": result.get("level", 1),
                    "prefix": result.get("prefix", ""),
                    "confidence": result.get("confidence", 0.8),
                    "reason": result.get("reason", ""),
                }

            # 统计
            h0 = sum(1 for r in results.values() if r["level"] == 0)
            h1 = sum(1 for r in results.values() if r["level"] == 1)
            h2 = sum(1 for r in results.values() if r["level"] == 2)
            prefixed = sum(1 for r in results.values() if r["prefix"])
            logger.info(
                "LLM 分析完成: %d 段落, ilvl=0:%d ilvl=1:%d ilvl=2:%d, 有前缀:%d",
                len(results), h0, h1, h2, prefixed,
            )
            return results

        except Exception as e:
            logger.warning("LLM 分析失败，将使用 fallback 规则: %s", e)
            return {}  # 返回空，触发 fallback

    # ── 参考文档模式 ──────────────────────────────────────────

    def _normalize_with_reference(self) -> None:  # pragma: no cover
        """从参考文档复制完整格式到测试文档"""
        assert self.ref_doc is not None and self.doc is not None

        # 1. 页面布局
        self._copy_page_layout()

        # 2. 编号定义
        self._copy_numbering_from_reference()

        # 3. 页眉页脚
        self._copy_header_footer()

        # 4. 获取参考文档的格式模板
        ref_body = self._get_ref_format("body")
        ref_h0 = self._get_ref_format("heading")
        ref_num_id = self._get_ref_numbering_num_id()

        # 5. 逐段落处理
        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            self._apply_reference_format(para, i, ref_body, ref_h0, ref_num_id)

        logger.debug("参考文档格式复制完成")

    def _get_ref_format(self, role: str) -> dict[str, Any]:  # pragma: no cover
        """从参考文档提取指定角色的格式模板"""
        assert self.ref_doc is not None
        ref_paras = [p for p in self.ref_doc.paragraphs if p.text.strip()]
        if not ref_paras:
            return {}

        target_ilvl = "0" if role == "heading" else "1"
        for p in ref_paras:
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    ilvl_el = numPr.find(qn("w:ilvl"))
                    if ilvl_el is not None and ilvl_el.get(qn("w:val")) == target_ilvl:
                        return self._extract_run_format(p)

        # fallback：取第一个或最后一个段落
        return self._extract_run_format(ref_paras[0] if role == "heading" else ref_paras[-1])

    def _extract_run_format(self, para: Any) -> dict[str, Any]:  # pragma: no cover
        """提取段落的 run 格式属性"""
        if not para.runs:
            return {}
        rPr = para.runs[0]._element.find(qn("w:rPr"))
        if rPr is None:
            return {}
        fmt: dict[str, Any] = {}
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                val = rFonts.get(qn(f"w:{attr}"))
                if val:
                    fmt[f"font_{attr}"] = val
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            fmt["sz"] = sz.get(qn("w:val"))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is not None:
            fmt["szCs"] = szCs.get(qn("w:val"))
        if rPr.find(qn("w:b")) is not None:
            fmt["bold"] = True
        return fmt

    def _get_ref_numbering_num_id(self) -> str:  # pragma: no cover
        """获取参考文档的主要编号 numId"""
        assert self.ref_doc is not None
        for para in self.ref_doc.paragraphs:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    numId_el = numPr.find(qn("w:numId"))
                    if numId_el is not None:
                        val = numId_el.get(qn("w:val"))
                        if val and val != "0":
                            return val  # type: ignore[no-any-return]
        return "1"

    # ── 段落格式应用 ─────────────────────────────────────────

    def _apply_reference_format(  # pragma: no cover
        self, para: Any, index: int, ref_body: dict, ref_h0: dict, num_id: str
    ) -> None:
        """对测试文档段落应用参考格式"""
        pPr = para._element.get_or_add_pPr()

        # 清除旧格式
        self._clear_format(pPr)

        # 行距 1.5 倍
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "360")
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)

        # 获取层级（LLM 优先，fallback 到规则）
        level = self._get_level(para, index)

        # 缩进
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:firstLine"), "400" if level < 2 else "402")
        pPr.append(ind)

        # 编号
        numPr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(level))
        numPr.append(ilvl_el)
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), num_id)
        numPr.append(numId_el)
        pPr.append(numPr)

        # 剥离手动编号（LLM 返回的 prefix 或 fallback 规则）
        self._strip_prefix(para, index)

        # 应用参考文档的 run 格式
        ref_fmt = ref_h0 if level == 0 else ref_body
        self._apply_run_format(para, ref_fmt, override_bold=(level == 0))

    def _apply_run_format(  # pragma: no cover
        self, para: Any, ref_fmt: dict, override_bold: bool = False
    ) -> None:
        """将参考格式应用到段落的所有 run"""
        if not ref_fmt:
            return

        for run in para.runs:
            rPr = run._element.get_or_add_rPr()

            # 清除旧的格式属性
            for tag in ("w:rFonts", "w:sz", "w:szCs", "w:b"):
                old = rPr.find(qn(tag))
                if old is not None:
                    rPr.remove(old)

            # 字体
            if any(k.startswith("font_") for k in ref_fmt):
                rFonts = OxmlElement("w:rFonts")
                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                    key = f"font_{attr}"
                    if key in ref_fmt:
                        rFonts.set(qn(f"w:{attr}"), ref_fmt[key])
                rPr.insert(0, rFonts)

            # 字号
            if "sz" in ref_fmt:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), ref_fmt["sz"])
                rPr.append(sz)
            if "szCs" in ref_fmt:
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), ref_fmt["szCs"])
                rPr.append(szCs)

            # 加粗
            if override_bold or ref_fmt.get("bold"):
                rPr.append(OxmlElement("w:b"))

    # ── 层级判断（LLM + fallback） ───────────────────────────

    def _get_level(self, para: Any, index: int) -> int:  # pragma: no cover
        """获取段落层级：LLM 结果优先，fallback 到简单规则"""
        # 优先用 LLM 结果
        if index in self._llm_results:
            return self._llm_results[index]["level"]  # type: ignore[no-any-return]
        # fallback
        return self._fallback_classify(para)

    def _fallback_classify(self, para: Any) -> int:  # pragma: no cover
        """简单规则 fallback（LLM 不可用时）

        基本思路：
        - 短粗体段落 → 标题 (ilvl=0)
        - 以中文数字+顿号、括号数字开头 → 根据模式判断层级
        - 其他 → 正文 (ilvl=1)
        """
        text = para.text.strip()
        if not text:
            return 1

        # 检测粗体
        is_bold = False
        if para.runs:
            rPr = para.runs[0]._element.find(qn("w:rPr"))
            if rPr is not None and rPr.find(qn("w:b")) is not None:
                is_bold = True

        # 短粗体 → 标题
        if is_bold and len(text) < 30:
            return 0

        # 以"一、"~"十、"开头 → 标题
        if len(text) >= 2 and text[0] in "一二三四五六七八九十" and len(text) > 1 and text[1] == "、":
            return 0

        # 以"（一）"~"（十）"开头 → 标题
        if text.startswith("（") and len(text) >= 3 and text[2] == "）":
            inner = text[1]
            if inner in "一二三四五六七八九十":
                return 0

        # 以数字+顿号/点开头 → 子项
        if len(text) >= 2 and text[0].isdigit():
            if len(text) > 1 and text[1] in "、.":
                return 2

        # 默认正文
        return 1

    # ── 前缀剥离（LLM + fallback） ──────────────────────────

    def _strip_prefix(self, para: Any, index: int) -> None:  # pragma: no cover
        """剥离手动编号前缀：LLM 结果优先，fallback 到规则"""
        text = para.text
        if not text:
            return

        new_text = text
        stripped_text = text.strip()

        # 优先用 LLM 返回的 prefix
        if index in self._llm_results:
            llm_prefix = self._llm_results[index].get("prefix", "")
            if llm_prefix and stripped_text.startswith(llm_prefix):
                # 保留前导空白，去掉 prefix
                leading = text[: len(text) - len(text.lstrip())]
                new_text = leading + stripped_text[len(llm_prefix):]
                new_text = new_text.lstrip("、．.,，：:").lstrip()
                # 恢复前导空白
                new_text = leading + new_text[len(leading):] if new_text.startswith(leading) else new_text

        # LLM 没有 prefix 时，用简单规则 fallback
        if new_text == text:
            new_text = self._fallback_strip(text)

        if new_text == text or not new_text.strip():
            return

        # 更新段落文本（保留第一个 run 的格式）
        self._replace_para_text(para, new_text)

    def _fallback_strip(self, text: str) -> str:  # pragma: no cover
        """简单规则剥离手动编号（LLM 不可用时的 fallback）"""
        import re

        # 保留前导空白
        leading = text[: len(text) - len(text.lstrip())]
        core = text[len(leading):]

        patterns = [
            # 多级编号（优先）："1.2.", "2.3.", "1.2.3."
            re.compile(r'^\d+(?:\.\d+)+[、．.：:]?\s*'),
            # 括号中文数字："（一）", "(一)", "（一）、", "(一)、"
            re.compile(r'^[（(][一二三四五六七八九十]+[)）][、．.：:]?\s*'),
            # 中文数字+顿号："一、"
            re.compile(r'^[一二三四五六七八九十]+、\s*'),
            # 括号阿拉伯数字："(1)", "（1）"
            re.compile(r'^[（(]\d+[)）][、．.：:]?\s*'),
            # 单级数字编号："1、", "2.", "3．"
            re.compile(r'^\d+[、．.]\s*'),
        ]

        for pattern in patterns:
            match = pattern.match(core)
            if match:
                result = core[match.end():]
                # 清理残留标点
                result = result.lstrip("、．.,，：:").lstrip()
                if result:
                    return leading + result
        return text

    def _replace_para_text(self, para: Any, new_text: str) -> None:  # pragma: no cover
        """替换段落文本，保留格式"""
        if para.runs:
            first_run = para.runs[0]
            # 清除所有 run
            for run in list(para.runs):
                run._element.getparent().remove(run._element)
            # 创建新 run，保持原格式
            new_run_el = OxmlElement("w:r")
            rPr = first_run._element.find(qn("w:rPr"))
            if rPr is not None:
                from lxml import etree

                new_run_el.append(etree.fromstring(etree.tostring(rPr)))
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = new_text
            new_run_el.append(t)
            para._element.append(new_run_el)

    # ── 页面布局 ─────────────────────────────────────────────

    def _copy_page_layout(self) -> None:  # pragma: no cover
        """从参考文档复制页面布局"""
        assert self.ref_doc is not None and self.doc is not None
        ref_sec = self.ref_doc.sections[0]
        doc_sec = self.doc.sections[0]

        doc_sec.top_margin = ref_sec.top_margin
        doc_sec.bottom_margin = ref_sec.bottom_margin
        doc_sec.left_margin = ref_sec.left_margin
        doc_sec.right_margin = ref_sec.right_margin
        doc_sec.page_width = ref_sec.page_width
        doc_sec.page_height = ref_sec.page_height

        ref_pgMar = ref_sec._sectPr.find(qn("w:pgMar"))
        doc_pgMar = doc_sec._sectPr.find(qn("w:pgMar"))
        if ref_pgMar is not None and doc_pgMar is not None:
            for attr in ("header", "footer"):
                val = ref_pgMar.get(qn(f"w:{attr}"))
                if val:
                    doc_pgMar.set(qn(f"w:{attr}"), val)

        logger.debug("页面布局已复制")

    # ── 编号定义 ─────────────────────────────────────────────

    def _copy_numbering_from_reference(self) -> None:  # pragma: no cover
        """从参考文档复制完整的编号定义"""
        assert self.ref_doc is not None and self.doc is not None

        ref_numbering = None
        try:
            ref_numbering = self.ref_doc.part.numbering_part
        except (KeyError, NotImplementedError):
            logger.warning("参考文档无编号定义")
            return

        ref_elm = ref_numbering._element
        ref_abstracts = ref_elm.findall(qn("w:abstractNum"))
        ref_nums = ref_elm.findall(qn("w:num"))

        if not ref_abstracts or not ref_nums:
            logger.warning("参考文档编号定义为空")
            return

        try:
            doc_numbering = self.doc.part.numbering_part
            doc_elm = doc_numbering._element
        except (KeyError, NotImplementedError):
            doc_elm = self._create_numbering_part()

        # 清除旧定义
        for old in doc_elm.findall(qn("w:abstractNum")):
            doc_elm.remove(old)
        for old in doc_elm.findall(qn("w:num")):
            doc_elm.remove(old)

        # 复制
        from lxml import etree

        for abstract in ref_abstracts:
            doc_elm.insert(0, etree.fromstring(etree.tostring(abstract)))
        for num in ref_nums:
            doc_elm.append(etree.fromstring(etree.tostring(num)))

        logger.debug("编号定义已复制: %d abstractNum, %d num", len(ref_abstracts), len(ref_nums))

    def _create_numbering_part(self) -> Any:  # pragma: no cover
        """手动创建编号 part"""
        assert self.doc is not None
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part
        from lxml import etree

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        numbering_elm = etree.Element(qn("w:numbering"), nsmap=nsmap)
        numbering_xml = etree.tostring(
            numbering_elm, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        part_name = PackURI("/word/numbering.xml")
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"
        numbering_part = Part(part_name, content_type, numbering_xml, self.doc.part.package)
        self.doc.part.relate_to(numbering_part, RT.NUMBERING)
        self._numbering_part = numbering_part

        return numbering_elm

    # ── 页眉页脚 ─────────────────────────────────────────────

    def _copy_header_footer(self) -> None:  # pragma: no cover
        """从参考文档复制页眉页脚（含页码域代码）"""
        assert self.ref_doc is not None and self.doc is not None

        for ref_sec, doc_sec in zip(self.ref_doc.sections, self.doc.sections):
            # ── 页眉 ──
            try:
                ref_header = ref_sec.header
                if not ref_header.is_linked_to_previous:
                    doc_header = doc_sec.header
                    doc_header.is_linked_to_previous = False
                    for p in doc_header.paragraphs:
                        p._element.getparent().remove(p._element)
                    for ref_p in ref_header.paragraphs:
                        new_p = doc_header._element.makeelement(qn("w:p"), {})
                        from lxml import etree

                        pPr = ref_p._element.find(qn("w:pPr"))
                        new_p.append(
                            etree.fromstring(etree.tostring(pPr))
                            if pPr is not None
                            else OxmlElement("w:pPr")
                        )
                        for ref_run in ref_p.runs:
                            new_p.append(etree.fromstring(etree.tostring(ref_run._element)))
                        doc_header._element.append(new_p)
            except Exception as e:
                logger.debug("页眉复制跳过: %s", e)

            # ── 页脚 ──
            try:
                ref_footer = ref_sec.footer
                if not ref_footer.is_linked_to_previous:
                    doc_footer = doc_sec.footer
                    doc_footer.is_linked_to_previous = False

                    has_page_field = False
                    for p in ref_footer.paragraphs:
                        for instr in p._element.findall('.//' + qn("w:instrText")):
                            if instr.text and "PAGE" in instr.text:
                                has_page_field = True
                                break

                    if has_page_field:
                        self._clear_footer(doc_footer)
                        self._add_page_number_field(doc_footer, ref_footer)
                        logger.debug("页脚：页码域代码已复制")
                    else:
                        self._clear_footer(doc_footer)
                        logger.debug("页脚：已清空")
            except Exception as e:
                logger.debug("页脚复制跳过: %s", e)

    def _clear_footer(self, footer: Any) -> None:  # pragma: no cover
        """清空页脚内容"""
        for child in list(footer._element):
            if child.tag != qn("w:pPr"):
                footer._element.remove(child)
        if not footer.paragraphs:
            footer.add_paragraph()

    def _add_page_number_field(self, doc_footer: Any, ref_footer: Any) -> None:  # pragma: no cover
        """添加 PAGE 域代码到页脚"""
        ref_pPr = None
        ref_run_fmt = None
        if ref_footer.paragraphs:
            ref_p = ref_footer.paragraphs[0]
            ref_pPr = ref_p._element.find(qn("w:pPr"))
            if ref_p.runs:
                ref_run_fmt = ref_p.runs[0]._element.find(qn("w:rPr"))

        para = OxmlElement("w:p")

        if ref_pPr is not None:
            from lxml import etree

            para.append(etree.fromstring(etree.tostring(ref_pPr)))
        else:
            pPr = OxmlElement("w:pPr")
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            pPr.append(jc)
            para.append(pPr)

        def _make_rPr() -> Any:  # pragma: no cover
            if ref_run_fmt is not None:
                from lxml import etree

                return etree.fromstring(etree.tostring(ref_run_fmt))
            return None

        # fldChar begin
        run1 = OxmlElement("w:r")
        rPr1 = _make_rPr()
        if rPr1 is not None:
            run1.append(rPr1)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run1.append(fldChar1)
        para.append(run1)

        # instrText
        run2 = OxmlElement("w:r")
        rPr2 = _make_rPr()
        if rPr2 is not None:
            run2.append(rPr2)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2.append(instrText)
        para.append(run2)

        # fldChar separate
        run3 = OxmlElement("w:r")
        rPr3 = _make_rPr()
        if rPr3 is not None:
            run3.append(rPr3)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run3.append(fldChar2)
        para.append(run3)

        # 占位符
        run4 = OxmlElement("w:r")
        rPr4 = _make_rPr()
        if rPr4 is not None:
            run4.append(rPr4)
        t = OxmlElement("w:t")
        t.text = "1"
        run4.append(t)
        para.append(run4)

        # fldChar end
        run5 = OxmlElement("w:r")
        rPr5 = _make_rPr()
        if rPr5 is not None:
            run5.append(rPr5)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run5.append(fldChar3)
        para.append(run5)

        doc_footer._element.append(para)

    # ── 工具方法 ─────────────────────────────────────────────

    def _clear_format(self, pPr: Any) -> None:  # pragma: no cover
        """清除段落中的旧格式定义"""
        for tag in ("w:spacing", "w:ind", "w:jc", "w:numPr"):
            old = pPr.find(qn(tag))
            if old is not None:
                pPr.remove(old)

    # ── 默认模式（无参考文档） ────────────────────────────────

    def _normalize_default(self) -> None:  # pragma: no cover
        """无参考文档时使用内置默认格式"""
        assert self.doc is not None
        from docx.shared import Cm

        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        self._setup_default_numbering()
        num_id = "1"

        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            pPr = para._element.get_or_add_pPr()
            self._clear_format(pPr)

            # 行距
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:line"), "360")
            spacing.set(qn("w:lineRule"), "auto")
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            pPr.append(spacing)

            # 层级
            level = self._get_level(para, i)

            # 缩进
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:firstLine"), "400" if level < 2 else "402")
            pPr.append(ind)

            # 编号
            numPr = OxmlElement("w:numPr")
            ilvl_el = OxmlElement("w:ilvl")
            ilvl_el.set(qn("w:val"), str(level))
            numPr.append(ilvl_el)
            numId_el = OxmlElement("w:numId")
            numId_el.set(qn("w:val"), num_id)
            numPr.append(numId_el)
            pPr.append(numPr)

            # 剥离前缀
            self._strip_prefix(para, i)

            # 默认字体
            for run in para.runs:
                rPr = run._element.get_or_add_rPr()
                for tag in ("w:rFonts", "w:sz", "w:szCs"):
                    old = rPr.find(qn(tag))
                    if old is not None:
                        rPr.remove(old)

                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "宋体")
                rFonts.set(qn("w:hAnsi"), "宋体")
                rFonts.set(qn("w:eastAsia"), "宋体")
                rPr.insert(0, rFonts)

                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "24")
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "24")
                rPr.append(szCs)

                if level == 0:
                    if rPr.find(qn("w:b")) is None:
                        rPr.append(OxmlElement("w:b"))

    def _setup_default_numbering(self) -> None:  # pragma: no cover
        """创建默认编号定义"""
        assert self.doc is not None

        try:
            numbering_elm = self.doc.part.numbering_part._element
        except (KeyError, NotImplementedError):
            numbering_elm = self._create_numbering_part()

        for old in numbering_elm.findall(qn("w:abstractNum")):
            numbering_elm.remove(old)
        for old in numbering_elm.findall(qn("w:num")):
            numbering_elm.remove(old)

        abstractNum = OxmlElement("w:abstractNum")
        abstractNum.set(qn("w:abstractNumId"), "0")

        levels = [
            ("0", "chineseCounting", "%1、", "400"),
            ("1", "decimal", "%2．", "400"),
            ("2", "decimal", "（%3）", "402"),
            ("3", "decimalEnclosedCircleChinese", "%4 ", "402"),
            ("4", "decimal", "%5）", "402"),
            ("5", "lowerLetter", "%6．", "402"),
            ("6", "lowerLetter", "%7）", "402"),
            ("7", "lowerRoman", "%8. ", "402"),
            ("8", "lowerRoman", "%9）", "402"),
        ]
        for ilvl, num_fmt, lvl_text, first_line in levels:
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), ilvl)
            lvl.set(qn("w:tentative"), "0")

            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)

            numFmt = OxmlElement("w:numFmt")
            numFmt.set(qn("w:val"), num_fmt)
            lvl.append(numFmt)

            suff = OxmlElement("w:suff")
            suff.set(qn("w:val"), "nothing")
            lvl.append(suff)

            lvlText = OxmlElement("w:lvlText")
            lvlText.set(qn("w:val"), lvl_text)
            lvl.append(lvlText)

            lvlJc = OxmlElement("w:lvlJc")
            lvlJc.set(qn("w:val"), "left")
            lvl.append(lvlJc)

            pPr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:firstLine"), first_line)
            pPr.append(ind)
            lvl.append(pPr)

            rPr = OxmlElement("w:rPr")
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:hint"), "eastAsia")
            rPr.append(rFonts)
            lvl.append(rPr)

            abstractNum.append(lvl)

        numbering_elm.insert(0, abstractNum)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstractNumRef = OxmlElement("w:abstractNumId")
        abstractNumRef.set(qn("w:val"), "0")
        num.append(abstractNumRef)
        numbering_elm.append(num)

        if hasattr(self, "_numbering_part"):
            from lxml import etree

            self._numbering_part._blob = etree.tostring(
                numbering_elm, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        logger.debug("默认编号定义已创建")
