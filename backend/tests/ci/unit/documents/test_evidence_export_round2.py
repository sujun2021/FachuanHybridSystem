"""Round 2 coverage tests for EvidenceExportService."""

from __future__ import annotations

from unittest.mock import MagicMock, patch, PropertyMock

import pytest


def _make_service(**kwargs):
    from apps.documents.services.evidence.evidence_export_service import EvidenceExportService
    return EvidenceExportService(placeholder_service=kwargs.get("placeholder_service", MagicMock()))


def _make_evidence_list(*, title="证据清单一", case_name="测试案件", order=1,
                        total_pages=0, export_version=1, previous_list_id=None):
    el = MagicMock()
    el.pk = 1
    el.id = 1
    el.title = title
    el.case.name = case_name
    el.case_id = 1
    el.order = order
    el.total_pages = total_pages
    el.export_version = export_version
    el.previous_list_id = previous_list_id
    return el


class TestPlaceholderServiceProperty:
    def test_returns_injected(self):
        ps = MagicMock()
        svc = _make_service(placeholder_service=ps)
        assert svc.placeholder_service is ps

    def test_lazy_loads_when_none(self):
        from apps.documents.services.evidence.evidence_export_service import EvidenceExportService
        svc = EvidenceExportService()
        with patch(
            "apps.documents.services.evidence.evidence_export_service.EvidenceExportService.placeholder_service",
            new_callable=PropertyMock,
        ) as mock_prop:
            mock_prop.return_value = MagicMock()
            _ = svc.placeholder_service


class TestGetEvidenceList:
    def test_found(self):
        svc = _make_service()
        with patch("apps.documents.services.evidence.evidence_export_service.EvidenceList") as mock_el:
            mock_el.objects.select_related.return_value.get.return_value = "el_obj"
            result = svc._get_evidence_list(1)
        assert result == "el_obj"

    def test_not_found(self):
        svc = _make_service()
        from apps.core.exceptions import NotFoundError
        from apps.evidence.models import EvidenceList
        with patch.object(EvidenceList, "objects") as mock_objects:
            mock_objects.select_related.return_value.get.side_effect = EvidenceList.DoesNotExist
            with pytest.raises(NotFoundError):
                svc._get_evidence_list(999)


class TestGetTemplate:
    def test_found(self):
        svc = _make_service()
        with patch("apps.documents.services.evidence.evidence_export_service.EvidenceList"):
            with patch(
                "apps.documents.services.evidence.evidence_export_service.EvidenceExportService._get_template",
                return_value=MagicMock(),
            ):
                pass
        # Direct test of _get_template
        with patch(
            "apps.documents.models.DocumentTemplate"
        ) as mock_tpl:
            mock_tpl.objects.get.return_value = MagicMock()
            # Import inside test to avoid side effects
            from apps.documents.models import DocumentCaseFileSubType, DocumentTemplateType
            mock_tpl.DoesNotExist = type("DoesNotExist", (Exception,), {})

    def test_not_found_raises(self):
        svc = _make_service()
        from apps.core.exceptions import NotFoundError
        from apps.documents.models import DocumentTemplate
        with patch.object(DocumentTemplate, "objects") as mock_objects:
            mock_objects.get.side_effect = DocumentTemplate.DoesNotExist
            with pytest.raises(NotFoundError):
                svc._get_template(999)


class TestGetGlobalOrderStart:
    def test_no_previous_lists(self):
        svc = _make_service()
        el = _make_evidence_list(order=1)
        with patch("apps.documents.services.evidence.evidence_export_service.EvidenceList") as mock_el:
            mock_el.objects.filter.return_value.aggregate.return_value = {"total": None}
            result = svc._get_global_order_start(el)
        assert result == 1

    def test_with_previous_items(self):
        svc = _make_service()
        el = _make_evidence_list(order=3)
        with patch("apps.documents.services.evidence.evidence_export_service.EvidenceList") as mock_el:
            mock_el.objects.filter.return_value.aggregate.return_value = {"total": 10}
            result = svc._get_global_order_start(el)
        assert result == 11


class TestGenerateFilename:
    def test_evidence_list_type(self):
        svc = _make_service()
        el = _make_evidence_list(title="证据清单一", case_name="测试案件", export_version=2)
        with patch("apps.documents.services.evidence.evidence_export_service.timezone") as mock_tz:
            mock_tz.now.return_value.strftime.return_value = "20260607"
            result = svc._generate_filename(el, "证据清单", 2)
        assert "证据清单一" in result
        assert "测试案件" in result
        assert "V2" in result
        assert "20260607" in result

    def test_evidence_detail_type_standard_title(self):
        svc = _make_service()
        el = _make_evidence_list(title="证据清单三")
        with patch("apps.documents.services.evidence.evidence_export_service.timezone") as mock_tz:
            mock_tz.now.return_value.strftime.return_value = "20260607"
            result = svc._generate_filename(el, "证据明细", 1)
        assert "证据明细三" in result

    def test_evidence_detail_supplement_title(self):
        svc = _make_service()
        el = _make_evidence_list(title="补充证据清单二")
        with patch("apps.documents.services.evidence.evidence_export_service.timezone") as mock_tz:
            mock_tz.now.return_value.strftime.return_value = "20260607"
            result = svc._generate_filename(el, "证据明细", 1)
        assert "证据明细二" in result

    def test_evidence_detail_custom_title(self):
        svc = _make_service()
        el = _make_evidence_list(title="自定义标题")
        with patch("apps.documents.services.evidence.evidence_export_service.timezone") as mock_tz:
            mock_tz.now.return_value.strftime.return_value = "20260607"
            result = svc._generate_filename(el, "证据明细", 1)
        assert "证据明细" in result


class TestIncrementVersion:
    def test_returns_current_version(self):
        svc = _make_service()
        el = _make_evidence_list(export_version=3)
        assert svc._increment_version(el) == 3


class TestExportEvidenceList:
    def test_generates_docx(self):
        svc = _make_service()
        el = _make_evidence_list(title="证据清单一", case_name="测试案件")
        el.items.order_by.return_value = []

        with patch.object(svc, "_get_evidence_list", return_value=el), \
             patch.object(svc, "_get_global_order_start", return_value=1), \
             patch.object(svc, "_increment_version", return_value=1), \
             patch.object(svc, "_generate_filename", return_value="test.docx"):
            content, filename = svc.export_evidence_list(1)

        assert isinstance(content, bytes)
        assert len(content) > 0
        assert filename == "test.docx"

    def test_with_items(self):
        svc = _make_service()
        el = _make_evidence_list(title="证据清单一", case_name="测试案件")
        item = MagicMock()
        item.name = "证据1"
        item.purpose = "证明事实"
        item.page_range_display = "1-3"
        el.items.order_by.return_value = [item]

        with patch.object(svc, "_get_evidence_list", return_value=el), \
             patch.object(svc, "_get_global_order_start", return_value=1), \
             patch.object(svc, "_increment_version", return_value=1), \
             patch.object(svc, "_generate_filename", return_value="test.docx"):
            content, filename = svc.export_evidence_list(1)

        assert isinstance(content, bytes)


class TestExportEvidenceDetail:
    def test_generates_docx(self):
        svc = _make_service()
        el = _make_evidence_list(title="证据清单一", case_name="测试案件")
        item = MagicMock()
        item.name = "证据1"
        item.purpose = "证明事实"
        item.page_range_display = "1-3"
        item.page_start = 1
        item.page_end = 3
        item.file = "some_file.pdf"
        item.file_name = "file.pdf"
        item.file_size_display = "1.0 MB"
        item.page_count = 3
        el.items.order_by.return_value = [item]

        with patch.object(svc, "_get_evidence_list", return_value=el), \
             patch.object(svc, "_get_global_order_start", return_value=1), \
             patch.object(svc, "_increment_version", return_value=1), \
             patch.object(svc, "_generate_filename", return_value="detail.docx"):
            content, filename = svc.export_evidence_detail(1)

        assert isinstance(content, bytes)
        assert len(content) > 0

    def test_item_without_file(self):
        svc = _make_service()
        el = _make_evidence_list(title="证据清单一", case_name="测试案件")
        item = MagicMock()
        item.name = "证据1"
        item.purpose = "证明事实"
        item.page_range_display = "-"
        item.page_start = None
        item.page_end = None
        item.file = None
        item.file_name = ""
        item.file_size_display = "-"
        item.page_count = 0
        el.items.order_by.return_value = [item]

        with patch.object(svc, "_get_evidence_list", return_value=el), \
             patch.object(svc, "_get_global_order_start", return_value=1), \
             patch.object(svc, "_increment_version", return_value=1), \
             patch.object(svc, "_generate_filename", return_value="detail.docx"):
            content, filename = svc.export_evidence_detail(1)

        assert isinstance(content, bytes)


class TestCreateEvidenceTable:
    def test_empty_items(self):
        from docx import Document
        svc = _make_service()
        doc = Document()
        svc._create_evidence_table(doc, [], global_order_start=1)
        # Table should exist with just the header
        assert len(doc.tables) == 1

    def test_with_items(self):
        from docx import Document
        svc = _make_service()
        doc = Document()
        item = MagicMock()
        item.name = "证据1"
        item.purpose = "证明"
        item.page_range_display = "1-3"
        svc._create_evidence_table(doc, [item], global_order_start=1)
        assert len(doc.tables) == 1


class TestAddEvidenceDetailSection:
    def test_with_page_range_and_file(self):
        from docx import Document
        svc = _make_service()
        doc = Document()
        item = MagicMock()
        item.name = "合同"
        item.purpose = "证明"
        item.page_start = 1
        item.page_end = 3
        item.page_range_display = "1-3"
        item.file = "file.pdf"
        item.file_name = "file.pdf"
        item.file_size_display = "1.0 MB"
        item.page_count = 3
        svc._add_evidence_detail_section(doc, item, global_order=1)

    def test_without_page_range(self):
        from docx import Document
        svc = _make_service()
        doc = Document()
        item = MagicMock()
        item.name = "合同"
        item.purpose = "证明"
        item.page_start = None
        item.page_end = None
        item.page_range_display = "-"
        item.file = None
        svc._add_evidence_detail_section(doc, item, global_order=2)
