"""Final coverage tests for evidence module."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from unittest.mock import MagicMock, Mock, patch

import pytest

from apps.evidence.services.export.evidence_export_service import EvidenceExportService
from apps.core.exceptions import NotFoundError, ValidationException


# ============================================================================
# EvidenceExportService init tests
# ============================================================================


class TestEvidenceExportServiceInit:
    def test_default_init(self):
        svc = EvidenceExportService()
        assert svc._placeholder_service is None

    def test_injected_init(self):
        mock_ps = MagicMock()
        svc = EvidenceExportService(placeholder_service=mock_ps)
        assert svc.placeholder_service is mock_ps


# ============================================================================
# _get_evidence_list tests
# ============================================================================


class TestGetEvidenceList:
    @pytest.mark.django_db
    def test_get_evidence_list_not_found(self):
        from apps.evidence.models import EvidenceList
        svc = EvidenceExportService()
        # Use a non-existent ID; the real model's DoesNotExist will be raised and caught
        with pytest.raises(NotFoundError):
            svc._get_evidence_list(999999999)

    @patch("apps.evidence.services.export.evidence_export_service.EvidenceList")
    def test_found(self, MockList):
        mock_el = MagicMock()
        MockList.objects.select_related.return_value.get.return_value = mock_el
        svc = EvidenceExportService()
        result = svc._get_evidence_list(1)
        assert result == mock_el


# ============================================================================
# _increment_version tests
# ============================================================================


class TestIncrementVersion:
    def test_returns_current_version(self):
        svc = EvidenceExportService()
        el = MagicMock()
        el.export_version = 3
        assert svc._increment_version(el) == 3


# ============================================================================
# _generate_filename tests
# ============================================================================


class TestGenerateFilename:
    @patch("apps.evidence.services.export.evidence_export_service.FilenameTemplateService")
    @patch("apps.evidence.services.export.evidence_export_service.timezone")
    def test_evidence_list_type(self, mock_tz, mock_fts):
        mock_tz.now.return_value.strftime.return_value = "20250601"
        mock_fts.render_generated_doc.return_value = "证据清单一(案件)V1_20250601"
        svc = EvidenceExportService()
        el = MagicMock()
        el.case.name = "案件"
        el.title = "证据清单一"
        result = svc._generate_filename(el, "证据清单", 1)
        assert result.endswith(".docx")

    @patch("apps.evidence.services.export.evidence_export_service.FilenameTemplateService")
    @patch("apps.evidence.services.export.evidence_export_service.timezone")
    def test_evidence_detail_type(self, mock_tz, mock_fts):
        mock_tz.now.return_value.strftime.return_value = "20250601"
        mock_fts.render_generated_doc.return_value = "证据明细一(案件)V1_20250601"
        svc = EvidenceExportService()
        el = MagicMock()
        el.case.name = "案件"
        el.title = "证据清单一"
        result = svc._generate_filename(el, "证据明细", 1)
        assert "证据明细" in result or result.endswith(".docx")

    @patch("apps.evidence.services.export.evidence_export_service.FilenameTemplateService")
    @patch("apps.evidence.services.export.evidence_export_service.timezone")
    def test_supplement_title(self, mock_tz, mock_fts):
        mock_tz.now.return_value.strftime.return_value = "20250601"
        mock_fts.render_generated_doc.return_value = "证据明细(案件)V1_20250601"
        svc = EvidenceExportService()
        el = MagicMock()
        el.case.name = "案件"
        el.title = "补充证据清单二"
        result = svc._generate_filename(el, "证据明细", 1)
        assert result.endswith(".docx")


# ============================================================================
# _add_evidence_detail_section tests
# ============================================================================


class TestAddEvidenceDetailSection:
    def test_adds_detail(self):
        from docx import Document
        svc = EvidenceExportService()
        doc = Document()
        item = MagicMock()
        item.direction = "plaintiff"
        item.get_direction_display.return_value = "原告"
        item.evidence_type = "contract"
        item.get_evidence_type_display.return_value = "合同"
        item.original_status = "original"
        item.get_original_status_display.return_value = "原件"
        item.purpose = "证明合同关系"
        item.page_start = 1
        item.page_end = 5
        item.page_range_display = "1-5"
        item.file = MagicMock()
        item.file_name = "contract.pdf"
        item.file_size_display = "1.2MB"
        item.page_count = 5
        item.name = "合同证据"
        svc._add_evidence_detail_section(doc, item, 1)
        assert len(doc.paragraphs) > 0

    def test_minimal_item(self):
        from docx import Document
        svc = EvidenceExportService()
        doc = Document()
        item = MagicMock()
        item.direction = None
        item.evidence_type = None
        item.original_status = None
        item.purpose = "证明内容"
        item.page_start = None
        item.page_end = None
        item.file = None
        item.name = "证据"
        svc._add_evidence_detail_section(doc, item, 1)
        assert len(doc.paragraphs) > 0


# ============================================================================
# export_evidence_list tests
# ============================================================================


class TestExportEvidenceList:
    @pytest.mark.django_db
    @patch("apps.evidence.services.export.evidence_export_service.EvidenceList")
    def test_export_creates_docx(self, MockList):
        svc = EvidenceExportService()
        mock_el = MagicMock()
        mock_el.title = "证据清单一"
        mock_el.case.name = "测试案件"
        mock_el.case_id = 1
        mock_el.order = 1
        mock_el.export_version = 1
        mock_el.items.order_by.return_value = []
        MockList.objects.select_related.return_value.get.return_value = mock_el
        MockList.objects.filter.return_value.aggregate.return_value = {"total": 0}

        content, filename = svc.export_evidence_list(1)
        assert isinstance(content, bytes)
        assert filename.endswith(".docx")


# ============================================================================
# export_evidence_detail tests
# ============================================================================


class TestExportEvidenceDetail:
    @pytest.mark.django_db
    @patch("apps.evidence.services.export.evidence_export_service.EvidenceList")
    def test_export_creates_docx(self, MockList):
        svc = EvidenceExportService()
        mock_el = MagicMock()
        mock_el.title = "证据清单一"
        mock_el.case.name = "测试案件"
        mock_el.case_id = 1
        mock_el.order = 1
        mock_el.export_version = 1
        mock_el.items.order_by.return_value = []
        MockList.objects.select_related.return_value.get.return_value = mock_el
        MockList.objects.filter.return_value.aggregate.return_value = {"total": 0}

        content, filename = svc.export_evidence_detail(1)
        assert isinstance(content, bytes)
        assert "证据明细" in filename or filename.endswith(".docx")


# ============================================================================
# export_zip tests
# ============================================================================


class TestExportZip:
    @pytest.mark.django_db
    @patch("apps.evidence.services.export.evidence_export_service.EvidenceList")
    def test_export_zip(self, MockList):
        svc = EvidenceExportService()
        mock_el = MagicMock()
        mock_el.title = "证据清单一"
        mock_el.case.name = "案件"
        mock_el.case_id = 1
        mock_el.order = 1
        mock_el.export_version = 1
        mock_el.items.order_by.return_value = []
        MockList.objects.select_related.return_value.get.return_value = mock_el
        MockList.objects.filter.return_value.aggregate.return_value = {"total": 0}

        content, filename = svc.export_zip(1)
        assert isinstance(content, bytes)
        assert filename.endswith(".zip")
        # Verify it's a valid zip
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            assert len(zf.namelist()) >= 1


# ============================================================================
# export_evidence_list_with_template tests
# ============================================================================


class TestExportWithTemplate:
    @pytest.mark.django_db
    @patch("apps.evidence.services.export.evidence_export_service.EvidenceList")
    def test_no_template_falls_back_to_default(self, MockList):
        svc = EvidenceExportService()
        mock_el = MagicMock()
        mock_el.title = "证据清单一"
        mock_el.case.name = "案件"
        mock_el.case_id = 1
        mock_el.order = 1
        mock_el.export_version = 1
        mock_el.items.order_by.return_value = []
        MockList.objects.select_related.return_value.get.return_value = mock_el
        MockList.objects.filter.return_value.aggregate.return_value = {"total": 0}

        content, filename = svc.export_evidence_list_with_template(1, template_id=None)
        assert isinstance(content, bytes)
        assert filename.endswith(".docx")
