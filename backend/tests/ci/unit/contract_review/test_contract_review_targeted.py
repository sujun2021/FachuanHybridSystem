"""Targeted tests for contract_review module to push coverage to 80%+."""
from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest


# ---------------------------------------------------------------------------
# format_normalizer/paragraph_classifier.py (0% coverage)
# ---------------------------------------------------------------------------


class TestParagraphClassifier:
    def setup_method(self):
        from apps.contract_review.services.format_normalizer.paragraph_classifier import ParagraphClassifier

        self.classifier = ParagraphClassifier()

    def test_classify_level0_service_content(self):
        level, reason = self.classifier.classify("服务内容")
        assert level == 0
        assert reason == "一级标题"

    def test_classify_level0_long_text_not_match(self):
        level, reason = self.classifier.classify("服务内容这是一个很长很长很长很长的文本描述这里")
        assert level != 0

    def test_classify_level1_number_prefix(self):
        level, reason = self.classifier.classify("1、甲方的义务")
        assert level == 1

    def test_classify_level1_chinese_number(self):
        level, reason = self.classifier.classify("一、合同条款")
        assert level == 1

    def test_classify_level1_has_party_keyword(self):
        level, reason = self.classifier.classify("乙方应当按照约定")
        assert level == 1

    def test_detect_level0_multiple_keywords(self):
        keywords = ["保密义务", "违约责任", "争议解决", "免责条款", "合同期限"]
        for kw in keywords:
            result = self.classifier._detect_level0(kw)
            assert result == 0, f"Expected level 0 for '{kw}'"

    def test_detect_level1_digit_with_keyword(self):
        result = self.classifier._detect_level1("1、甲方的义务", "")
        assert result == 1

    def test_detect_level1_chinese_number(self):
        result = self.classifier._detect_level1("一、合同总则", "")
        assert result == 1

    def test_detect_level2_dot_format(self):
        result = self.classifier._detect_level2("1. 安装要求", "")
        assert result == 2

    def test_detect_level2_with_keyword(self):
        result = self.classifier._detect_level2("需要维修和检测", "")
        assert result == 2

    def test_detect_level2_no_match(self):
        result = self.classifier._detect_level2("xyz", "")
        assert result == -1


# ---------------------------------------------------------------------------
# formatting/page_numbering.py (19% coverage)
# ---------------------------------------------------------------------------


class TestPageNumbering:
    def test_standardize(self):
        from docx import Document

        from apps.contract_review.services.formatting.page_numbering import PageNumbering

        doc = Document()
        doc.add_paragraph("Test content")
        pn = PageNumbering()
        pn.standardize(doc)
        section = doc.sections[-1]
        footer = section.footer
        footer_xml = footer._element.xml
        assert "w:fldChar" in footer_xml

    def test_add_page_number_footer(self):
        from docx import Document

        from apps.contract_review.services.formatting.page_numbering import PageNumbering

        doc = Document()
        pn = PageNumbering()
        pn._add_page_number_footer(doc)
        section = doc.sections[-1]
        footer = section.footer
        footer_xml = footer._element.xml
        assert "w:fldChar" in footer_xml


# ---------------------------------------------------------------------------
# extraction/content_extractor.py (45% coverage)
# ---------------------------------------------------------------------------


class TestContentExtractor:
    def test_extract_paragraphs(self, tmp_path):
        from docx import Document

        from apps.contract_review.services.extraction.content_extractor import ContentExtractor

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("")
        doc.add_paragraph("Third paragraph")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        extractor = ContentExtractor()
        result = extractor.extract_paragraphs(f)
        assert "First paragraph" in result
        assert "Third paragraph" in result
        assert "" not in result

    def test_extract_with_mapping(self, tmp_path):
        from docx import Document

        from apps.contract_review.services.extraction.content_extractor import ContentExtractor

        doc = Document()
        doc.add_paragraph("Para 0")
        doc.add_paragraph("Para 1")
        doc.add_paragraph("")
        doc.add_paragraph("Para 3")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        extractor = ContentExtractor()
        result = extractor.extract_with_mapping(f)
        assert result.paragraphs == ["Para 0", "Para 1", "Para 3"]
        assert result.index_map == [0, 1, 3]

    def test_extract_empty_doc(self, tmp_path):
        from docx import Document

        from apps.contract_review.services.extraction.content_extractor import ContentExtractor

        doc = Document()
        doc.add_paragraph("")
        f = tmp_path / "empty.docx"
        doc.save(str(f))

        extractor = ContentExtractor()
        with pytest.raises(Exception):
            extractor.extract_paragraphs(f)


# ---------------------------------------------------------------------------
# extraction/title_extractor.py (79% coverage)
# ---------------------------------------------------------------------------


class TestTitleExtractor:
    def test_parse_title_from_filename(self):
        from apps.contract_review.services.extraction.title_extractor import TitleExtractor

        result = TitleExtractor.parse_title_from_filename("某某合同.docx")
        assert isinstance(result, str)

    def test_parse_title_from_filename_no_ext(self):
        from apps.contract_review.services.extraction.title_extractor import TitleExtractor

        result = TitleExtractor.parse_title_from_filename("contract")
        assert isinstance(result, str)


# ---------------------------------------------------------------------------
# tasks.py (49% coverage)
# ---------------------------------------------------------------------------


class TestContractReviewTasks:
    @patch("apps.contract_review.tasks.ReviewTask")
    def test_cleanup_old_files_no_tasks(self, mock_model):
        from apps.contract_review.tasks import cleanup_old_files

        mock_model.objects.filter.return_value.exclude.return_value = []
        result = cleanup_old_files(days=30)
        assert result["upload_files"] == 0
        assert result["output_files"] == 0

    @patch("apps.contract_review.tasks.ReviewTask")
    @patch("apps.contract_review.repositories.review_task_repository.ReviewTaskRepository")
    def test_cleanup_old_files_with_tasks(self, mock_repo_cls, mock_model):
        from apps.contract_review.tasks import cleanup_old_files

        task = SimpleNamespace(id=1, original_file=None, output_file=None)
        mock_qs = MagicMock()
        mock_qs.__iter__ = MagicMock(return_value=iter([task]))
        mock_model.objects.filter.return_value.exclude.return_value = mock_qs
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo

        result = cleanup_old_files(days=30)
        assert result["upload_files"] == 0
        assert result["output_files"] == 0
        assert result["tasks"] == 1


# ---------------------------------------------------------------------------
# models/format_normalize_detail.py (62% coverage)
# ---------------------------------------------------------------------------


class TestFormatNormalizeDetail:
    def test_model_meta(self):
        from apps.contract_review.models.format_normalize_detail import FormatNormalizeDetail

        assert FormatNormalizeDetail._meta.verbose_name == "格式调整详情"
        assert FormatNormalizeDetail._meta.ordering == ["-created_at"]


# ---------------------------------------------------------------------------
# services/contract_format_service.py (74% coverage)
# ---------------------------------------------------------------------------


class TestContractFormatService:
    def test_import_service(self):
        from apps.contract_review.services.contract_format_service import ContractFormatService

        assert ContractFormatService is not None
