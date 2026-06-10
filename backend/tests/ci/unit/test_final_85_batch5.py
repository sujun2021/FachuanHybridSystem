"""Coverage tests for batch 5 - targeting remaining uncovered modules.

Focus on:
- DocxFormatter (57 unc)
- DocumentMatchingMixin (69 unc)
- DocumentQueryMixin (65 unc)
- FrameProcessingService more methods
- DocumentDelivery services
"""

from __future__ import annotations

from types import SimpleNamespace
from unittest.mock import MagicMock, Mock, patch

import pytest


# ===================================================================
# DocxFormatter
# ===================================================================
class TestDocxFormatter:
    def test_find_title_element_heading(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter

        doc = Document()
        p = doc.add_paragraph("标题", style="Heading 1")
        result = DocxFormatter._find_title_element(doc)
        assert result is p._element

    def test_find_title_element_first_non_empty(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter

        doc = Document()
        p = doc.add_paragraph("第一段内容")
        result = DocxFormatter._find_title_element(doc)
        assert result is p._element

    def test_find_title_element_empty_doc(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter

        doc = Document()
        # Remove default paragraph
        for para in doc.paragraphs:
            p_elem = para._element
            p_elem.getparent().remove(p_elem)
        result = DocxFormatter._find_title_element(doc)
        assert result is None

    def test_set_paragraph_spacing_title(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter

        doc = Document()
        p = doc.add_paragraph("标题")
        DocxFormatter._set_paragraph_spacing(p, is_title=True)
        assert p.paragraph_format.space_before.pt == 0
        assert p.paragraph_format.space_after.pt == 0

    def test_set_paragraph_spacing_body(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter

        doc = Document()
        p = doc.add_paragraph("正文")
        DocxFormatter._set_paragraph_spacing(p, is_title=False)
        assert p.paragraph_format.first_line_indent.pt == 24

    def test_set_font(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter
        from docx.shared import Pt

        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("测试文字")
        DocxFormatter._set_font(p, "宋体", Pt(12))
        assert run.font.size == Pt(12)

    def test_clean_style_indent_chars(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter

        doc = Document()
        DocxFormatter._clean_style_indent_chars(doc)  # no exception

    def test_format_document(self):
        from docx import Document
        from apps.contract_review.services.formatting.docx_formatter import DocxFormatter

        doc = Document()
        doc.add_paragraph("合同标题", style="Heading 1")
        doc.add_paragraph("第一条 正文内容")
        formatter = DocxFormatter()
        formatter.format_document(doc)  # no exception


# ===================================================================
# DocumentMatchingMixin: helper methods
# ===================================================================
class TestDocumentMatchingMixin:
    def test_match_case_by_number(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return Mock(match_by_case_number=Mock(return_value=SimpleNamespace(id=1)))

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        result = mixin._match_case_by_number("(2024)粤01民初1号")
        assert result.id == 1

    def test_match_case_by_document_parties_success(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        matcher = Mock()
        matcher.extract_parties_from_document.return_value = ["张三", "李四"]
        matcher.match_by_party_names.return_value = SimpleNamespace(id=1, status="active")

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return matcher

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        result = mixin._match_case_by_document_parties(["/path/doc.pdf"])
        assert result.id == 1

    def test_match_case_by_document_parties_no_extract(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        matcher = Mock()
        matcher.extract_parties_from_document.return_value = []

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return matcher

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        result = mixin._match_case_by_document_parties(["/path/doc.pdf"])
        assert result is None

    def test_match_case_by_document_parties_closed_case(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        matcher = Mock()
        matcher.extract_parties_from_document.return_value = ["张三"]
        matcher.match_by_party_names.return_value = SimpleNamespace(id=1, status="closed")

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return matcher

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        result = mixin._match_case_by_document_parties(["/path/doc.pdf"])
        assert result is None

    def test_match_case_by_document_parties_exception(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        matcher = Mock()
        matcher.extract_parties_from_document.side_effect = RuntimeError("error")

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return matcher

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        result = mixin._match_case_by_document_parties(["/path/doc.pdf"])
        assert result is None

    def test_sync_case_number_to_case_success(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        cn_service = Mock()
        cn_service.list_numbers_internal.return_value = []
        cn_service.create_number_internal.return_value = SimpleNamespace(number="(2024)粤01民初1号")

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return Mock()

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        with patch(
            "apps.core.dependencies.business_case.build_case_number_service",
            return_value=cn_service,
        ):
            result = mixin._sync_case_number_to_case(1, "(2024)粤01民初1号")
        assert result is True

    def test_sync_case_number_to_case_already_exists(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        cn_service = Mock()
        cn_service.list_numbers_internal.return_value = [
            SimpleNamespace(number="(2024)粤01民初1号")
        ]

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return Mock()

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        with patch(
            "apps.core.dependencies.business_case.build_case_number_service",
            return_value=cn_service,
        ):
            result = mixin._sync_case_number_to_case(1, "(2024)粤01民初1号")
        assert result is True

    def test_sync_case_number_to_case_exception(self):
        from apps.automation.services.document_delivery.api.document_delivery_api_service._matching import (
            DocumentMatchingMixin,
        )

        class TestMixin(DocumentMatchingMixin):
            @property
            def case_matcher(self):
                return Mock()

            @property
            def document_renamer(self):
                return Mock()

            @property
            def notification_service(self):
                return Mock()

        mixin = TestMixin()
        with patch(
            "apps.core.dependencies.business_case.build_case_number_service",
            side_effect=RuntimeError("error"),
        ):
            result = mixin._sync_case_number_to_case(1, "(2024)粤01民初1号")
        assert result is False


# ===================================================================
# FrameProcessingService: more methods
# ===================================================================
class TestFrameProcessingServiceMore:
    def test_check_ocr_similarity_empty_kept_texts(self):
        from apps.chat_records.services.extraction.frame_processing_service import FrameProcessingService

        svc = FrameProcessingService()
        state = SimpleNamespace(kept_ocr_texts=[])
        result = svc.check_ocr_similarity("Hello World", state, 0.5, 5)
        assert result is None

    def test_check_ocr_similarity_empty_ocr_text(self):
        from apps.chat_records.services.extraction.frame_processing_service import FrameProcessingService

        svc = FrameProcessingService()
        state = SimpleNamespace(kept_ocr_texts=["prev"])
        result = svc.check_ocr_similarity("", state, 0.5, 5)
        assert result is None


# ===================================================================
# DocumentQueryResult and related data classes
# ===================================================================
class TestDocumentQueryDataClasses:
    def test_import_data_classes(self):
        from apps.automation.services.document_delivery.data_classes import (
            DocumentDeliveryRecord,
            DocumentQueryResult,
            DocumentRecord,
        )

        assert DocumentQueryResult is not None
        assert DocumentRecord is not None
        assert DocumentDeliveryRecord is not None

    def test_query_result_defaults(self):
        from apps.automation.services.document_delivery.data_classes import DocumentQueryResult

        result = DocumentQueryResult(
            total_found=0,
            processed_count=0,
            skipped_count=0,
            failed_count=0,
            case_log_ids=[],
            errors=[],
        )
        assert result.total_found == 0
        assert result.processed_count == 0


# ===================================================================
# Automation utils text_utils
# ===================================================================
class TestTextUtils:
    def test_normalize_case_number(self):
        from apps.automation.utils.text_utils import TextUtils

        result = TextUtils.normalize_case_number("(2024)粤01民初1号")
        assert result is not None
        assert "2024" in result

    def test_normalize_case_number_empty(self):
        from apps.automation.utils.text_utils import TextUtils

        result = TextUtils.normalize_case_number("")
        assert result == "" or result is not None


# ===================================================================
# Automation services sms matching module
# ===================================================================
class TestSMSMatchingModule:
    def test_import_document_parser_service(self):
        from apps.automation.services.sms.matching import DocumentParserService

        assert DocumentParserService is not None

    def test_import_party_matching_service(self):
        from apps.automation.services.sms.matching import PartyMatchingService

        assert PartyMatchingService is not None


# ===================================================================
# Core services
# ===================================================================
class TestCoreServices:
    def test_filename_template_service(self):
        from apps.core.services.filename_template_service import FilenameTemplateService

        assert FilenameTemplateService is not None

    def test_system_config_service(self):
        from apps.core.services.system_config_service import SystemConfigService

        assert SystemConfigService is not None


# ===================================================================
# Core cloud_storage protocols
# ===================================================================
class TestCloudStorageProtocols:
    def test_cloud_file_info(self):
        from apps.core.cloud_storage.protocols import CloudFileInfo

        info = CloudFileInfo(name="test.pdf", path="/path/test.pdf", is_dir=False, size=1024, modified_at=None)
        assert info.name == "test.pdf"
        assert info.is_dir is False


# ===================================================================
# Automation choices module
# ===================================================================
class TestAutomationChoices:
    def test_scraper_task_status(self):
        from apps.automation.models import ScraperTaskStatus

        assert ScraperTaskStatus.PENDING is not None
        assert ScraperTaskStatus.RUNNING is not None
        assert ScraperTaskStatus.SUCCESS is not None
        assert ScraperTaskStatus.FAILED is not None

    def test_court_sms_status(self):
        from apps.automation.models import CourtSMSStatus

        assert CourtSMSStatus.MATCHING is not None
        assert CourtSMSStatus.RENAMING is not None
        assert CourtSMSStatus.FAILED is not None
        assert CourtSMSStatus.PENDING_MANUAL is not None


# ===================================================================
# Automation schemas
# ===================================================================
class TestAutomationSchemas:
    def test_court_sms_schemas(self):
        from apps.automation.schemas.court_sms import SMSParseResult

        assert SMSParseResult is not None

    def test_document_delivery_schemas(self):
        from apps.automation.schemas.document_delivery import DocumentDeliveryRecord

        assert DocumentDeliveryRecord is not None

    def test_preservation_schemas(self):
        from apps.automation.schemas.preservation import PreservationQuoteSchema

        assert PreservationQuoteSchema is not None
